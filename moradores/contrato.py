"""Gera o Contrato de Locação Residencial por Temporada em .docx,
preenchido com os dados do morador e os dados específicos da locação."""
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _p(doc, texto="", *, bold=False, align=None, size=11, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.bold = bold
    run.font.size = Pt(size)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)
    return p


def _clausula(doc, texto):
    _p(doc, texto, align="justify", space_after=8)


def gerar_contrato_docx(morador, dados):
    """morador: instância de Morador.
    dados: dict com chaves opcionais:
        endereco_locatario, cep_locatario, tel_fixo,
        endereco_imovel, numero_quarto, num_habitantes, num_adultos, num_criancas,
        valor, forma_pagamento, data_inicio, data_fim, data_contrato,
        estado_imovel, qualificacao_locador
    """
    doc = Document()

    # estilo base
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    _p(doc, "Contrato de Locação Residencial por Temporada", bold=True, align="center", size=14, space_after=12)

    _clausula(doc,
        "Os signatários, que contratam nas qualidades indicadas neste contrato, têm entre si, "
        "ajustada a presente locação residencial para temporada, A PRAZO CERTO, mediante as "
        "seguintes cláusulas e condições a seguir, e condições legais pertinentes que "
        "voluntariamente outorgam:")

    locador = dados.get("qualificacao_locador") or "(QUALIFICAÇÃO COMPLETA)"
    _clausula(doc, f"1 – LOCADOR: {locador}")

    # Dados do locatário (vêm do morador)
    endereco = dados.get("endereco_locatario") or "______________"
    cep = dados.get("cep_locatario") or "______________"
    tel_fixo = dados.get("tel_fixo") or "__________________"
    celular = morador.whatsapp or morador.telefone or "_____________"
    _clausula(doc,
        f"2 – LOCATÁRIO: {morador.nome}, CPF/MF: {morador.cpf or '_______________'}, "
        f"RG {morador.rg or '_______________'} SSP/SP, residente e domiciliado à {endereco} - "
        f"Cep: {cep}. Contatos: Tel Fixo: {tel_fixo}  Celular: {celular}.")

    quarto = dados.get("numero_quarto") or (morador.quarto.nome if morador.quarto else "______")
    endereco_imovel = dados.get("endereco_imovel") or "(Endereço do imóvel)"
    _clausula(doc, f"3 - OBJETO DA LOCAÇÃO: {endereco_imovel} - Quarto: {quarto}.")

    estado = dados.get("estado_imovel") or (
        "O Locatário receberá o imóvel em perfeito estado de apresentação, limpo, com todos os "
        "sistemas de água, energia e sanitário em funcionamento normal, com Móveis, Utensílios e "
        "Eletrodomésticos (fora roupa de cama e banho), banheiro coletivo, tudo de acordo com "
        "vistoria a ser efetuada no recebimento das chaves, comprometendo-se a devolver no mesmo estado.")
    _clausula(doc, f"3.1 - ESTADO DO IMÓVEL: {estado}")

    num_hab = dados.get("num_habitantes") or "_____"
    num_ad = dados.get("num_adultos") or "_________"
    num_cri = dados.get("num_criancas") or "_________"
    _clausula(doc,
        f"4 - FIM A QUE SE DESTINA: Fins residenciais a prazo certo (temporada), com número "
        f"limitado de habitantes ({num_hab}) - sendo {num_ad} adultos e {num_cri} crianças.")

    valor = dados.get("valor") or (f"{morador.valor_contratado:.2f}".replace(".", ",") if morador.valor_contratado else "xxxxxx")
    forma = dados.get("forma_pagamento") or "xxxxxxxx"
    _clausula(doc, f"5 - VALOR DA LOCAÇÃO: R$ {valor}. O aluguel é o indicado neste contrato devendo ser pago {forma}.")
    _clausula(doc,
        "Deverá ser entregue pelo locatário, antecipadamente ao recebimento das chaves, cópia "
        "dos documentos e endereço utilizados neste contrato, bem como cópia do seu RG com CPF ou CNH.")

    di = dados.get("data_inicio") or "xxxxxx"
    df = dados.get("data_fim") or "xxxxxxx"
    _clausula(doc,
        f"6 - PRAZO DE LOCAÇÃO: De {di} até {df}, ocasião em que o imóvel deverá ser desocupado "
        f"e em ordem para vistoria de saída.")

    _clausula(doc, "7 - OBRIGAÇÕES GERAIS: O locatário, no recebimento das chaves, procederá vistoria do imóvel e ficará obrigado a:")
    _clausula(doc, "- Manter o objeto da locação no mais perfeito estado de conservação e limpo, para assim o restituir ao Locador no mesmo estado recebido.")
    _clausula(doc, "- Não fazer instalação, adaptação ou benfeitoria, sem prévia autorização por escrito do locador;")
    _clausula(doc, "- Não transferir este contrato, não sublocar, não ceder, emprestar ou aumentar o número de habitantes sem o prévio conhecimento do locador, sob qualquer pretexto e de igual forma, não alterar a destinação da locação, não constituindo o decurso do tempo, por si só, na demora do Locador em reprimir a infração, assentimento à mesma.")

    _clausula(doc, "9 - RESCISÃO CONTRATUAL: A infração das obrigações consignadas na cláusula sétima, sem prejuízo de qualquer outra prevista em lei, por parte do locatário, é considerada como de natureza grave, acarretando a rescisão contratual, com o consequente despejo e obrigatoriedade de imediata satisfação dos consectários contratuais e legais.")
    _clausula(doc, "- Em se tratando de um contrato por temporada o locador não será obrigado em hipótese alguma a devolver o valor pago adiantado ou quitado integralmente.")
    _clausula(doc, "10 – INDENIZAÇÃO: Em caso de dano de qualquer espécie ao imóvel ou objetos, por responsabilidade do locatário, será cobrado, em valor de mercado a reposição/reparação a título de indenização.")
    _clausula(doc, "11 – RENOVAÇÃO: É direito do locador a não renovação se assim o achar conveniente.")
    _clausula(doc, "12 - VANTAGENS LEGAIS SUPERVENIENTES: A locação estará sempre sujeita ao Regime do Código Civil Brasileiro e ficando assegurado ao locador todos os direitos e vantagens conferidas pela legislação que vier a ser promulgada durante locação.")
    _clausula(doc, "13 - CLÁUSULA PENAL: O locador e o locatário obrigam-se a respeitar o presente contrato em todas as suas cláusulas e condições, incorrendo a parte que infringir qualquer disposição contratual ou legal na multa igual ao valor de aluguel deste contrato, que será sempre paga integralmente qualquer que seja o tempo contratual decorrido, inclusive se verificada a prorrogação da vigência da locação.")
    _clausula(doc, "- O pagamento da multa não obsta a rescisão do contrato por parte inocente, caso lhe convier.")
    _clausula(doc, "- As partes contratantes elegem o foro da situação do imóvel, quaisquer que sejam os seus domicílios e, por estarem justos e contratados, assinam o presente instrumento em duas vias de igual teor, na presença das testemunhas que igualmente assinam.")

    data_contrato = dados.get("data_contrato") or "_____ de __________ de 2026"
    _p(doc, f"São Caetano do Sul, {data_contrato}.", space_after=18)

    _p(doc, "Locador: _________________________________________", space_after=14)
    _p(doc, f"Locatário: {morador.nome}", space_after=2)
    _p(doc, "Assinatura: _________________________________________", space_after=14)
    _p(doc, "Testemunha 1: _____________________________________", space_after=8)
    _p(doc, "Testemunha 2: _____________________________________")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
